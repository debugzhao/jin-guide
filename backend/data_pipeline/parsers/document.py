from __future__ import annotations

import re
from html.parser import HTMLParser
from pathlib import Path

from data_pipeline.records import DocumentChunkRecord, PolicyRuleRecord, Provenance
from data_pipeline.text_encoding import decode_html_bytes


KEY_SECTION_TERMS = (
    "录取",
    "调剂",
    "选科",
    "体检",
    "外语",
    "单科",
    "学费",
    "中外合作",
    "转专业",
    "专业介绍",
    "志愿",
    "投档",
)

# 按 document_type 分级的 (预算字符数, overlap字符数)。charter 条款长、逐条列举，
# 需要更大窗口保留完整语义；major_intro/policy 语义单元更短，小窗口检索更精准。
# 数值仍以字符数计量（保持跟历史生产数据同一套量纲，避免额外引入 token 估算误差）。
_CHUNK_BUDGET: dict[str, tuple[int, int]] = {
    "charter": (1400, 200),
    "policy": (900, 150),
    "transfer_policy": (900, 150),
    "major_intro": (900, 120),
}
_DEFAULT_CHUNK_BUDGET = (1200, 150)

# 导航/菜单类噪声段落的判定：真实招生页面里"首页/学院介绍/联系我们/网站导航"这类
# 菜单项通常连续密集出现（一次抓取往往会拿到十几个相邻菜单项），而真正的短标题
# （如"第四章 招生录取"、"十一、录取规则"）在正文里是孤立出现、后面紧跟长段正文。
# 用"连续 N 个短且无终止标点的段落"这个结构特征识别导航块，而不是维护一份不断
# 增长、容易漏检的菜单词表——每换一个学校官网模板，词表就要重新补一批。
_NAV_LINE_MAX_CHARS = 12
_NAV_RUN_MIN_LENGTH = 3
_TERMINAL_PUNCTUATION = ("。", "！", "？", "：", ":")


class _TextExtractor(HTMLParser):
    BLOCK_TAGS = {"p", "div", "li", "h1", "h2", "h3", "h4", "tr", "br"}

    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self._ignored = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag.lower() in {"script", "style", "noscript"}:
            self._ignored += 1
        elif tag.lower() in self.BLOCK_TAGS:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in {"script", "style", "noscript"}:
            self._ignored = max(0, self._ignored - 1)
        elif tag.lower() in self.BLOCK_TAGS:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._ignored:
            self.parts.append(data)


def extract_document_text(path: str | Path) -> str:
    file_path = Path(path)
    if file_path.suffix.lower() in {".html", ".htm"}:
        parser = _TextExtractor()
        parser.feed(decode_html_bytes(file_path.read_bytes()))
        text = "".join(parser.parts)
    elif file_path.suffix.lower() == ".txt":
        # 人工从图片/微信公众号文章转录的正文（`*.manual.txt`），已经是纯文本，
        # 不需要走任何格式还原，直接读取即可。
        text = file_path.read_text(encoding="utf-8")
    elif file_path.suffix.lower() == ".pdf":
        try:
            import pdfplumber
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("pdfplumber is required to extract PDF text") from exc
        with pdfplumber.open(file_path) as document:
            text = "\n".join(page.extract_text() or "" for page in document.pages)
    elif file_path.suffix.lower() == ".docx":
        # 浙江省招生政策通知附件是 .docx（江苏是 .pdf），官方 word 排版里正文段落和
        # 表格是分开的两种元素，逐段读 paragraphs 会漏掉表格里的内容（例如某些年份
        # 把批次控制线放在表格里），所以两者都读，按文档内出现顺序拼接。
        import docx

        document = docx.Document(str(file_path))
        parts: list[str] = []
        for element in document.element.body:
            tag = element.tag.rsplit("}", 1)[-1]
            if tag == "p":
                paragraph = next(
                    (p for p in document.paragraphs if p._element is element), None
                )
                if paragraph is not None and paragraph.text.strip():
                    parts.append(paragraph.text)
            elif tag == "tbl":
                table = next((t for t in document.tables if t._element is element), None)
                if table is not None:
                    for row in table.rows:
                        parts.append(" ".join(cell.text.strip() for cell in row.cells))
        text = "\n".join(parts)
    else:
        raise ValueError(f"unsupported document format: {file_path.suffix}")
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


_VHTML_PATTERN = re.compile(r"vHtml:\s*'((?:[^'\\]|\\.)*)'")


def extract_westlake_embedded_html_text(path: str | Path) -> str:
    """西湖大学招生动态详情页正文不在真实DOM里——页面骨架是内联`<script>new
    AdmissionEventsDetail("#app", {...vHtml:'<div>...(JSON式转义)...</div>'})`
    把正文整段塞进JS对象字面量的字符串属性里（已用真实页面curl+Playwright双重
    验证：DOM侧`#app`容器为空，纯httpx完全拿不到正文；但正文其实原样躺在下载
    到的原始HTML字节里，是"藏在script里"不是"真的需要JS渲染"，此前误判为必须
    走浏览器，实际只需要从原始文本里挖出这段转义字符串再反转义+去标签）。
    """
    file_path = Path(path)
    raw = decode_html_bytes(file_path.read_bytes())
    match = _VHTML_PATTERN.search(raw)
    if not match:
        return ""
    escaped = match.group(1)
    unescaped = (
        escaped.replace('\\"', '"')
        .replace("\\/", "/")
        .replace("\\n", "\n")
        .replace("\\'", "'")
    )
    parser = _TextExtractor()
    parser.feed(unescaped)
    text = "".join(parser.parts)
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def _strip_navigation_noise(paragraphs: list[str]) -> list[str]:
    """剔除导航/菜单类噪声段落。

    两步处理：①去掉连续重复的段落（同一导航栏在页面里常常原样出现两遍）；
    ②识别"连续 N 个短且无终止标点"的段落块，视为导航菜单整体剔除（见上方
    `_NAV_LINE_MAX_CHARS`/`_NAV_RUN_MIN_LENGTH` 的判定依据）。
    """
    deduped: list[str] = []
    for paragraph in paragraphs:
        if deduped and deduped[-1] == paragraph:
            continue
        deduped.append(paragraph)

    def _is_nav_like(paragraph: str) -> bool:
        return len(paragraph) <= _NAV_LINE_MAX_CHARS and not paragraph.endswith(_TERMINAL_PUNCTUATION)

    keep = [True] * len(deduped)
    i = 0
    while i < len(deduped):
        if not _is_nav_like(deduped[i]):
            i += 1
            continue
        j = i
        while j < len(deduped) and _is_nav_like(deduped[j]):
            j += 1
        if j - i >= _NAV_RUN_MIN_LENGTH:
            for k in range(i, j):
                keep[k] = False
        i = j
    return [paragraph for paragraph, kept in zip(deduped, keep) if kept]


def _select_paragraphs(paragraphs: list[str], document_type: str) -> list[str]:
    """policy/charter/transfer_policy 按关键词过滤段落，但保留命中段落的前后各一段。

    纯粹按关键词命中与否二选一会把条款正文之间的过渡句/子项一并连根拔起——
    实测某高校章程因此丢了近一半条款（第1/2/3/5/6/7/9/10/21/23/25/27条等
    完整消失，不是切分边界问题，是根本没进候选池）。保留紧邻的前后段落能捞回
    大部分"夹在两条命中条款中间"的短间隔，同时仍然丢弃真正大段无关的行政性
    文字（离最近一次关键词命中有 2 段以上距离的内容）。
    """
    if document_type not in {"policy", "charter", "transfer_policy"}:
        return paragraphs
    matched = [any(term in paragraph for term in KEY_SECTION_TERMS) for paragraph in paragraphs]
    keep = [
        matched[i]
        or (i > 0 and matched[i - 1])
        or (i < len(matched) - 1 and matched[i + 1])
        for i in range(len(paragraphs))
    ]
    return [paragraph for paragraph, kept in zip(paragraphs, keep) if kept]


_SECTION_BOUNDARY_PATTERN = re.compile(r"^第[一二三四五六七八九十百零〇0-9]+[条章]")
# 只要当前 chunk 已经攒够预算的一半，遇到新的"第N条/第N章"就提前切一刀，哪怕还没
# 到字符上限——纯粹按字符数贪心拼接会把物理上相邻但主题无关的条款也拼进同一个
# chunk（实测东华大学章程"体检色觉异常要求"和"学费标准"紧挨着排在一起，硬拼进
# 同一个 chunk 会稀释语义、拖累 rerank 命中率），优先在条款边界切能让每个 chunk
# 更贴近一个独立的语义单元。
_SECTION_BOUNDARY_SOFT_MIN_RATIO = 0.5


def _tail_overlap(text: str, overlap_chars: int) -> str:
    """取上一个 chunk 结尾的一小段作为下一个 chunk 的开头（overlap）。

    尽量从句子边界之后开始截取，避免 overlap 本身就是从半句话开始。
    """
    if overlap_chars <= 0 or not text:
        return ""
    tail = text[-overlap_chars:]
    boundary = re.search(r"[。！？\n]", tail)
    return tail[boundary.end():] if boundary else tail


def split_into_chunks(
    text: str,
    *,
    document_type: str,
    max_chars: int | None = None,
    overlap_chars: int | None = None,
) -> list[str]:
    """纯文本切分逻辑：导航噪声过滤 → 关键词过滤（保留上下文）→ 分级预算 + overlap 贪心拼接。

    不依赖 Provenance，供 `chunk_document`（在线采集）和离线重切分脚本共用。
    """
    default_max, default_overlap = _CHUNK_BUDGET.get(document_type, _DEFAULT_CHUNK_BUDGET)
    budget_chars = max_chars if max_chars is not None else default_max
    budget_overlap = overlap_chars if overlap_chars is not None else default_overlap

    paragraphs = [part.strip() for part in re.split(r"\n+", text) if part.strip()]
    paragraphs = _strip_navigation_noise(paragraphs)
    selected = _select_paragraphs(paragraphs, document_type)

    chunks: list[str] = []
    current = ""
    for paragraph in selected:
        would_overflow = current and len(current) + len(paragraph) + 1 > budget_chars
        at_section_boundary = (
            current
            and len(current) >= budget_chars * _SECTION_BOUNDARY_SOFT_MIN_RATIO
            and _SECTION_BOUNDARY_PATTERN.match(paragraph)
        )
        if current and (would_overflow or at_section_boundary):
            chunks.append(current)
            overlap = _tail_overlap(current, budget_overlap)
            current = f"{overlap}\n{paragraph}".strip() if overlap else paragraph
        else:
            current = f"{current}\n{paragraph}".strip() if current else paragraph
    if current:
        chunks.append(current)
    return chunks


def chunk_document(
    text: str,
    *,
    document_type: str,
    provenance: Provenance,
    province: str,
    university_code: str | None = None,
    max_chars: int | None = None,
    overlap_chars: int | None = None,
) -> list[DocumentChunkRecord]:
    chunks = split_into_chunks(
        text,
        document_type=document_type,
        max_chars=max_chars,
        overlap_chars=overlap_chars,
    )
    return [
        DocumentChunkRecord(
            province=province,
            year=provenance.year,
            document_type=document_type,
            university_code=university_code,
            chunk_index=index,
            content=content,
            provenance=provenance,
        )
        for index, content in enumerate(chunks)
    ]


def extract_policy_rule(text: str, *, provenance: Provenance, province: str) -> PolicyRuleRecord:
    compact = re.sub(r"\s+", "", text)
    mode = "parallel" if "平行志愿" in compact else "unknown"
    max_match = re.search(r"(?:设置|填报|可填报)(\d{1,3})个院校专业组志愿", compact)
    if max_match is None:
        # 浙江真实政策原文（2026年实施方案已验证）用"专业平行志愿"，数字和"志愿"之间
        # 没有"院校/专业"这类单位名（"考生每次可填报不超过80个志愿"），跟江苏"设置N个
        # 院校专业组志愿"结构不同。必须锁定"不超过N个志愿"且"个"后直接接"志愿"二字，
        # 否则会误命中同一段里"1个志愿单位"（志愿单位说明句，不是数量上限）或"不超过
        # 6个专业志愿"（提前录取院校志愿的专业数子上限，不是本志愿单位的总数上限）
        max_match = re.search(r"不超过(\d{1,3})个志愿(?!单位|专业|院校)", compact)
    adjustment_allowed = None
    if "服从专业调剂" in compact or "专业调剂" in compact:
        adjustment_allowed = True
    if "不进行专业调剂" in compact or "不服从专业调剂" in compact:
        adjustment_allowed = False

    tie_match = re.search(r"([^。]{0,80}(?:同分|投档分相同)[^。]{0,180}。)", text)
    filing_match = re.search(r"([^。]{0,80}(?:投档原则|投档规则)[^。]{0,220}。)", text)
    return PolicyRuleRecord(
        province=province,
        year=provenance.year,
        volunteer_mode=mode,
        max_volunteers=int(max_match.group(1)) if max_match else None,
        adjustment_allowed=adjustment_allowed,
        filing_rule=filing_match.group(1).strip() if filing_match else None,
        tie_break_rule=tie_match.group(1).strip() if tie_match else None,
        provenance=provenance,
    )
